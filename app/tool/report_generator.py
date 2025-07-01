import asyncio
import json
import os
from typing import List, Optional, Dict, Any
from pydantic import Field
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.tool.base import BaseTool, ToolResult
from app.llm import LLM


class ReportGenerator(BaseTool):
    name: str = "report_generator"
    description: str = """A specialized tool for generating comprehensive reports.

    This tool can generate detailed reports by breaking them into manageable chunks
    and then combining them into a complete document. It's designed to work around
    token limitations of local models.

    Use this when you need to create:
    - Business reports
    - Research papers
    - Technical documentation
    - Analysis reports
    - Summary reports

    The tool will automatically save the report as a Word document (.docx) file.
    """

    parameters: dict = {
        "type": "object",
        "properties": {
            "action": {
                "type": "string",
                "enum": ["generate_full_report", "generate_outline", "generate_section", "combine_sections"],
                "description": "The action to perform"
            },
            "topic": {
                "type": "string",
                "description": "The main topic or title of the report"
            },
            "requirements": {
                "type": "string",
                "description": "Specific requirements or instructions for the report"
            },
            "section_title": {
                "type": "string",
                "description": "Title of the section to generate (for generate_section action)"
            },
            "outline": {
                "type": "string",
                "description": "The report outline (for generate_section action)"
            },
            "output_format": {
                "type": "string",
                "enum": ["txt", "docx"],
                "default": "docx",
                "description": "Output format: txt for text, docx for Word document"
            }
        },
        "required": ["action"]
    }

    llm: LLM = Field(default_factory=LLM)
    sections: Dict[str, str] = Field(default_factory=dict, exclude=True)

    async def execute(self, **kwargs) -> ToolResult:
        action = kwargs.get("action")

        if action == "generate_full_report":
            return await self._generate_full_report(**kwargs)
        elif action == "generate_outline":
            return await self._generate_outline(**kwargs)
        elif action == "generate_section":
            return await self._generate_section(**kwargs)
        elif action == "combine_sections":
            return await self._combine_sections(**kwargs)
        else:
            return ToolResult(
                success=False,
                content=f"Unknown action: {action}"
            )

    async def _generate_full_report(self, **kwargs) -> ToolResult:
        """Generate a complete report in one go (for shorter reports)."""
        topic = kwargs.get("topic", "General Report")
        requirements = kwargs.get("requirements", "")
        output_format = kwargs.get("output_format", "docx")

        prompt = f"""请写一份关于"{topic}"的完整报告。

要求：
{requirements}

请按照以下结构组织报告：
1. 执行摘要
2. 引言
3. 主要内容（分2-3个小节）
4. 结论
5. 建议

请确保报告内容详实、逻辑清晰、结构完整。总字数控制在3000-5000字之间。"""

        try:
            response = await self.llm.agenerate(prompt)

            if output_format == "docx":
                # 保存为Word文件
                file_path = await self._save_as_word(topic, response, requirements)
                return ToolResult(
                    success=True,
                    content=f"报告已生成并保存为Word文件：{file_path}\n\n报告内容预览：\n{response[:500]}..."
                )
            else:
                return ToolResult(
                    success=True,
                    content=f"报告生成完成：\n\n{response}"
                )

        except Exception as e:
            return ToolResult(
                success=False,
                content=f"生成报告时出错：{str(e)}"
            )

    async def _generate_outline(self, **kwargs) -> ToolResult:
        """Generate a detailed outline for a long report."""
        topic = kwargs.get("topic", "General Report")
        requirements = kwargs.get("requirements", "")

        prompt = f"""请为关于"{topic}"的报告生成一个详细的大纲。

要求：
{requirements}

请生成一个包含以下内容的大纲：
1. 主标题和副标题
2. 每个章节的详细子标题
3. 每个章节的主要内容要点
4. 建议的字数分配

格式要求：
- 使用数字编号（1, 1.1, 1.1.1等）
- 每个章节包含3-5个子章节
- 总章节数控制在5-8个
- 包含执行摘要、引言、主要内容、结论、建议等必要部分"""

        try:
            response = await self.llm.agenerate(prompt)
            return ToolResult(
                success=True,
                content=f"报告大纲已生成：\n\n{response}"
            )
        except Exception as e:
            return ToolResult(
                success=False,
                content=f"生成大纲时出错：{str(e)}"
            )

    async def _generate_section(self, **kwargs) -> ToolResult:
        """Generate a specific section of the report."""
        section_title = kwargs.get("section_title", "")
        outline = kwargs.get("outline", "")

        if not section_title:
            return ToolResult(
                success=False,
                content="缺少section_title参数"
            )

        prompt = f"""请根据以下大纲，详细撰写"{section_title}"章节的内容。

报告大纲：
{outline}

要求：
1. 内容要详实、专业
2. 逻辑清晰，结构合理
3. 字数控制在800-1200字
4. 使用适当的标题和段落
5. 如果有数据或案例，请提供具体信息"""

        try:
            response = await self.llm.agenerate(prompt)

            # 保存章节内容
            self.sections[section_title] = response

            return ToolResult(
                success=True,
                content=f"章节'{section_title}'已生成：\n\n{response}"
            )
        except Exception as e:
            return ToolResult(
                success=False,
                content=f"生成章节时出错：{str(e)}"
            )

    async def _combine_sections(self, **kwargs) -> ToolResult:
        """Combine all generated sections into a complete report."""
        topic = kwargs.get("topic", "General Report")
        requirements = kwargs.get("requirements", "")
        output_format = kwargs.get("output_format", "docx")

        if not self.sections:
            return ToolResult(
                success=False,
                content="没有可合并的章节内容"
            )

        # 组合所有章节
        full_report = f"# {topic}\n\n"
        full_report += f"## 报告要求\n{requirements}\n\n"
        full_report += "---\n\n"

        for section_title, content in self.sections.items():
            full_report += f"## {section_title}\n\n{content}\n\n"

        # 添加总结
        full_report += "---\n\n"
        full_report += "## 报告总结\n\n"
        full_report += "本报告已完成所有章节的撰写，内容涵盖了相关主题的各个方面。"

        if output_format == "docx":
            # 保存为Word文件
            file_path = await self._save_as_word(topic, full_report, requirements)
            return ToolResult(
                success=True,
                content=f"报告已合并并保存为Word文件：{file_path}\n\n报告总字数：约{len(full_report)}字"
            )
        else:
            return ToolResult(
                success=True,
                content=f"报告合并完成：\n\n{full_report}"
            )

    async def _save_as_word(self, topic: str, content: str, requirements: str) -> str:
        """Save the report content as a Word document."""
        # 创建reports目录
        reports_dir = "reports"
        if not os.path.exists(reports_dir):
            os.makedirs(reports_dir)

        # 生成文件名
        import datetime
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_topic = "".join(c for c in topic if c.isalnum() or c in (' ', '-', '_')).rstrip()
        filename = f"{safe_topic}_{timestamp}.docx"
        file_path = os.path.join(reports_dir, filename)

        # 创建Word文档
        doc = Document()

        # 设置中文字体
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 添加标题
        title = doc.add_heading(topic, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加报告要求
        if requirements:
            doc.add_heading('报告要求', level=1)
            doc.add_paragraph(requirements)
            doc.add_paragraph()  # 空行

        # 处理内容
        lines = content.split('\n')
        current_heading_level = 0

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 检测标题
            if line.startswith('#'):
                heading_level = len(line) - len(line.lstrip('#'))
                heading_text = line.lstrip('#').strip()

                if heading_level == 1:
                    doc.add_heading(heading_text, level=1)
                elif heading_level == 2:
                    doc.add_heading(heading_text, level=2)
                elif heading_level == 3:
                    doc.add_heading(heading_text, level=3)
                else:
                    doc.add_heading(heading_text, level=4)

                current_heading_level = heading_level
            else:
                # 普通段落
                if line.startswith('- ') or line.startswith('* '):
                    # 列表项
                    p = doc.add_paragraph()
                    p.add_run('• ').bold = True
                    p.add_run(line[2:])
                else:
                    # 普通段落
                    doc.add_paragraph(line)

        # 保存文档
        doc.save(file_path)

        return file_path
